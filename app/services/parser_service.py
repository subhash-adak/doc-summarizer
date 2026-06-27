import io
from typing import Callable

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from app.core.exceptions import DocumentParseError, UnsupportedFileTypeError
from app.core.logging import get_logger
from app.models.document import DriveFile, ParsedDocument, SupportedMimeType

logger = get_logger(__name__)


# Individual parsers 

def _parse_pdf(content: bytes) -> tuple[str, int, list[str]]:
    """
    Extract text from a PDF using PyMuPDF.
    Returns (text, page_count, warnings).
    Falls back gracefully if a page has no selectable text (scanned PDF).
    """
    errors: list[str] = []
    pages: list[str] = []

    with fitz.open(stream=content, filetype="pdf") as doc:
        page_count = len(doc)
        for i, page in enumerate(doc):
            try:
                text = page.get_text("text")
                pages.append(text)
            except Exception as exc:
                errors.append(f"Page {i + 1}: {exc}")

    full_text = "\n".join(pages).strip()

    if not full_text and not errors:
        errors.append(
            "No selectable text found. Document may be a scanned image PDF. "
            "OCR is not currently supported."
        )

    return full_text, page_count, errors


def _parse_docx(content: bytes) -> tuple[str, None, list[str]]:
    """
    Extract text from a DOCX file using python-docx.
    Includes text from tables and headers.
    """
    errors: list[str] = []
    parts: list[str] = []

    try:
        doc = DocxDocument(io.BytesIO(content))

        # Body paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

    except Exception as exc:
        errors.append(f"DOCX parse error: {exc}")

    return "\n".join(parts).strip(), None, errors


def _parse_txt(content: bytes) -> tuple[str, None, list[str]]:
    """
    Decode plain-text file. Attempts UTF-8 then falls back to latin-1.
    """
    errors: list[str] = []
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        text = content.decode("latin-1")
        errors.append("File decoded with latin-1 fallback (not valid UTF-8).")
    return text.strip(), None, errors


# Dispatcher 

_PARSER_MAP: dict[SupportedMimeType, Callable] = {
    SupportedMimeType.PDF: _parse_pdf,
    SupportedMimeType.DOCX: _parse_docx,
    SupportedMimeType.TXT: _parse_txt,
    SupportedMimeType.MD: _parse_txt,
    SupportedMimeType.CSV: _parse_txt,
}


class ParserService:
    """
    Stateless document parser.
    Dispatches to the correct parser based on MIME type.
    """

    def parse(self, file: DriveFile, content: bytes) -> ParsedDocument:
        """
        Extract text from document bytes.

        Args:
            file:    DriveFile metadata (used for logging and the returned model)
            content: Raw file bytes downloaded from Drive

        Returns:
            ParsedDocument with extracted text (may have extraction_errors populated)

        Raises:
            UnsupportedFileTypeError: if the MIME type has no registered parser
        """
        parser_fn = _PARSER_MAP.get(file.mime_type)
        if parser_fn is None:
            raise UnsupportedFileTypeError(file.mime_type)

        logger.info(
            "parsing_document",
            file_id=file.id,
            file_name=file.name,
            mime_type=file.mime_type,
        )

        try:
            text, page_count, errors = parser_fn(content)
        except Exception as exc:
            logger.error(
                "parser_unexpected_error",
                file_id=file.id,
                file_name=file.name,
                error=str(exc),
            )
            raise DocumentParseError(
                f"Failed to parse '{file.name}': {exc}"
            ) from exc

        if errors:
            logger.warning(
                "parser_warnings",
                file_id=file.id,
                file_name=file.name,
                warnings=errors,
            )

        logger.info(
            "parsing_complete",
            file_id=file.id,
            char_count=len(text),
            page_count=page_count,
        )

        return ParsedDocument(
            file=file,
            raw_text=text,
            page_count=page_count,
            extraction_errors=errors,
        )
